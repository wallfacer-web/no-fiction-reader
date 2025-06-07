#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
英文非虚构图书阅读辅助软件 (English Non-Fiction Reading Assistant)
一个基于AI的英文非虚构图书阅读辅助工具，帮助中文英语专业学生更好地理解和学习英文非虚构类文本。

Author: Toby LUO@ZHKU (903098625@qq.com)
Copyright (c) 2025 Toby LUO@ZHKU (903098625@qq.com)
License: MIT License

基于《How to Read Non-Fiction English Books for Chinese English Majors》研究文献开发

GitHub: https://github.com/wallfacer-web/no-fiction-reader
"""

import gradio as gr
import requests
import json
import re
from typing import List, Dict, Any, Tuple
import time
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import logging
import math
from collections import Counter
import nltk
from nltk.tokenize import word_tokenize, sent_tokenize
from nltk.corpus import stopwords
import sqlite3
from datetime import datetime

# 设置日志
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# 下载必要的NLTK数据
try:
    nltk.data.find('tokenizers/punkt')
except LookupError:
    nltk.download('punkt')

try:
    nltk.data.find('corpora/stopwords')
except LookupError:
    nltk.download('stopwords')

class VocabularyDatabase:
    """词汇数据库管理"""
    
    def __init__(self, db_path: str = "vocabulary.db"):
        self.db_path = db_path
        self.init_database()
    
    def init_database(self):
        """初始化数据库"""
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        
        # 创建词汇表
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS vocabulary (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                word TEXT UNIQUE,
                definition TEXT,
                word_family TEXT,
                frequency_level INTEGER,
                learned_count INTEGER DEFAULT 0,
                first_seen DATE,
                last_reviewed DATE
            )
        ''')
        
        # 创建学习记录表
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS learning_progress (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                session_date DATE,
                words_learned INTEGER,
                paragraphs_processed INTEGER,
                reading_time INTEGER
            )
        ''')
        
        conn.commit()
        conn.close()
    
    def add_word(self, word: str, definition: str, word_family: str = "", frequency_level: int = 5):
        """添加单词到数据库"""
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        
        try:
            cursor.execute('''
                INSERT OR REPLACE INTO vocabulary 
                (word, definition, word_family, frequency_level, first_seen)
                VALUES (?, ?, ?, ?, ?)
            ''', (word.lower(), definition, word_family, frequency_level, datetime.now().date()))
            conn.commit()
        except Exception as e:
            logger.error(f"添加单词失败: {e}")
        finally:
            conn.close()
    
    def get_learned_words(self) -> List[str]:
        """获取已学习的单词列表"""
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        
        cursor.execute('SELECT word FROM vocabulary WHERE learned_count > 0')
        words = [row[0] for row in cursor.fetchall()]
        
        conn.close()
        return words

class TextDifficultyAnalyzer:
    """文本难度分析器 - 专门针对非虚构文本"""
    
    def __init__(self):
        # 基础常用词汇表（模拟前3000个最常用英语单词）
        self.common_words = self._load_basic_words()
        # 学术和非虚构文本常见词汇
        self.academic_words = self._load_academic_words()
    
    def _load_basic_words(self) -> set:
        """加载基础词汇表"""
        basic_words = [
            'the', 'be', 'to', 'of', 'and', 'a', 'in', 'that', 'have',
            'i', 'it', 'for', 'not', 'on', 'with', 'he', 'as', 'you',
            'do', 'at', 'this', 'but', 'his', 'by', 'from', 'they',
            'she', 'or', 'an', 'will', 'my', 'one', 'all', 'would',
            'there', 'their', 'what', 'so', 'up', 'out', 'if', 'about',
            'who', 'get', 'which', 'go', 'me', 'when', 'make', 'can',
            'like', 'time', 'no', 'just', 'him', 'know', 'take', 'people',
            'into', 'year', 'your', 'good', 'some', 'could', 'them', 'see',
            'other', 'than', 'then', 'now', 'look', 'only', 'come', 'its',
            'over', 'think', 'also', 'back', 'after', 'use', 'two', 'how',
            'our', 'work', 'first', 'well', 'way', 'even', 'new', 'want',
            'because', 'any', 'these', 'give', 'day', 'most', 'us', 'was',
            'been', 'said', 'each', 'which', 'she', 'do', 'how', 'their',
            'if', 'will', 'up', 'other', 'about', 'out', 'many', 'then',
            'them', 'these', 'so', 'some', 'her', 'would', 'make', 'like',
            'into', 'him', 'has', 'more', 'go', 'no', 'way', 'could', 'my',
            'than', 'first', 'water', 'been', 'call', 'who', 'its', 'now',
            'find', 'long', 'down', 'day', 'did', 'get', 'come', 'made',
            'may', 'part', 'over', 'new', 'sound', 'take', 'only', 'little',
            'work', 'know', 'place', 'year', 'live', 'me', 'back', 'give',
            'most', 'very', 'after', 'thing', 'our', 'name', 'good', 'sentence',
            'man', 'think', 'say', 'great', 'where', 'help', 'through', 'much',
            'before', 'line', 'right', 'too', 'mean', 'old', 'any', 'same',
            'tell', 'boy', 'follow', 'came', 'want', 'show', 'also', 'around',
            'form', 'three', 'small', 'set', 'put', 'end', 'why', 'again',
            'turn', 'here', 'off', 'went', 'old', 'number', 'great', 'tell',
            'men', 'say', 'small', 'every', 'found', 'still', 'between',
            'mane', 'should', 'home', 'big', 'give', 'air', 'line', 'set',
            'own', 'under', 'read', 'last', 'never', 'us', 'left', 'end',
            'along', 'while', 'might', 'next', 'sound', 'below', 'saw',
            'something', 'thought', 'both', 'few', 'those', 'always', 'looked',
            'show', 'large', 'often', 'together', 'asked', 'house', 'world',
            'going', 'want', 'school', 'important', 'until', 'without', 'form',
            'black', 'white', 'words', 'students', 'during', 'started', 'include',
            'young', 'book', 'example', 'took', 'being', 'different', 'state',
            'never', 'became', 'between', 'high', 'really', 'something', 'most',
            'another', 'much', 'family', 'own', 'out', 'leave', 'put', 'old',
            'while', 'mean', 'on', 'keep', 'student', 'why', 'let', 'great',
            'same', 'big', 'group', 'begin', 'seem', 'country', 'help', 'talk',
            'where', 'turn', 'problem', 'every', 'start', 'hand', 'might',
            'american', 'show', 'part', 'about', 'against', 'place', 'over',
            'such', 'again', 'few', 'case', 'most', 'week', 'company', 'where',
            'system', 'each', 'right', 'program', 'hear', 'so', 'question',
            'during', 'work', 'play', 'government', 'run', 'small', 'number',
            'off', 'always', 'move', 'like', 'night', 'live', 'mr', 'point',
            'believe', 'hold', 'today', 'bring', 'happen', 'next', 'without',
            'before', 'large', 'all', 'million', 'must', 'home', 'under', 'water',
            'room', 'write', 'mother', 'area', 'national', 'money', 'story',
            'young', 'fact', 'month', 'different', 'lot', 'right', 'study',
            'book', 'eye', 'job', 'word', 'though', 'business', 'issue', 'side',
            'kind', 'four', 'head', 'far', 'black', 'long', 'both', 'little',
            'house', 'yes', 'after', 'since', 'long', 'provide', 'service',
            'around', 'friend', 'important', 'father', 'sit', 'away', 'until',
            'power', 'hour', 'game', 'often', 'yet', 'line', 'political', 'end',
            'among', 'ever', 'stand', 'bad', 'lose', 'however', 'member', 'pay',
            'law', 'meet', 'car', 'city', 'almost', 'include', 'continue',
            'set', 'later', 'community', 'much', 'name', 'five', 'once', 'white',
            'least', 'president', 'learn', 'real', 'change', 'team', 'minute',
            'best', 'several', 'idea', 'kid', 'body', 'information', 'back',
            'parent', 'face', 'others', 'level', 'office', 'door', 'health',
            'person', 'art', 'war', 'history', 'party', 'within', 'grow',
            'result', 'open', 'change', 'morning', 'walk', 'reason', 'low',
            'win', 'research', 'girl', 'guy', 'early', 'food', 'before', 'moment',
            'himself', 'air', 'teacher', 'force', 'offer'
        ]
        return set(basic_words)
    
    def _load_academic_words(self) -> set:
        """加载学术和非虚构文本常用词汇"""
        academic_words = [
            'analysis', 'research', 'study', 'evidence', 'data', 'theory', 'concept',
            'argument', 'hypothesis', 'methodology', 'conclusion', 'discussion',
            'interpretation', 'significance', 'implication', 'perspective', 'framework',
            'approach', 'strategy', 'principle', 'factor', 'element', 'aspect',
            'dimension', 'variable', 'criterion', 'parameter', 'characteristic',
            'phenomenon', 'process', 'structure', 'function', 'relationship',
            'correlation', 'comparison', 'contrast', 'similarity', 'difference',
            'category', 'classification', 'definition', 'explanation', 'description',
            'evaluation', 'assessment', 'measurement', 'observation', 'investigation'
        ]
        return set(academic_words)
    
    def analyze_text_difficulty(self, text: str) -> Dict[str, Any]:
        """分析文本难度 - 专门针对非虚构文本"""
        # 简单的词汇分析
        words = re.findall(r'\b[a-zA-Z]+\b', text.lower())
        total_words = len(words)
        unique_words = len(set(words))
        
        # 句子分析
        sentences = re.split(r'[.!?]+', text)
        sentences = [s.strip() for s in sentences if s.strip()]
        
        # 计算常用词比例
        common_word_count = sum(1 for word in words if word in self.common_words)
        common_word_ratio = common_word_count / total_words if total_words > 0 else 0
        
        # 计算学术词汇比例
        academic_word_count = sum(1 for word in words if word in self.academic_words)
        academic_word_ratio = academic_word_count / total_words if total_words > 0 else 0
        
        # 计算平均句长
        avg_sentence_length = total_words / len(sentences) if sentences else 0
        
        # 识别难词和专业术语
        difficult_words = [word for word in set(words) 
                         if word not in self.common_words and len(word) > 3]
        technical_terms = [word for word in set(words) 
                         if word in self.academic_words]
        
        # 识别文本特征（标题、列表等）
        text_features = self._identify_text_features(text)
        
        # 计算难度评分 (1-10, 10最难) - 针对非虚构文本调整
        difficulty_score = self._calculate_nonfiction_difficulty_score(
            common_word_ratio, academic_word_ratio, avg_sentence_length, 
            len(difficult_words), unique_words, text_features
        )
        
        return {
            'total_words': total_words,
            'unique_words': unique_words,
            'common_word_ratio': common_word_ratio,
            'academic_word_ratio': academic_word_ratio,
            'avg_sentence_length': avg_sentence_length,
            'difficult_words': difficult_words[:15],
            'technical_terms': technical_terms[:10],
            'text_features': text_features,
            'difficulty_score': difficulty_score,
            'reading_level': self._get_nonfiction_reading_level(difficulty_score),
            'estimated_reading_time': self._estimate_nonfiction_reading_time(total_words),
            'vocabulary_coverage': common_word_ratio * 100,
            'academic_density': academic_word_ratio * 100
        }
    
    def _identify_text_features(self, text: str) -> Dict[str, int]:
        """识别非虚构文本特征"""
        features = {
            'headings': len(re.findall(r'^[A-Z][A-Za-z\s]*:?$', text, re.MULTILINE)),
            'numbered_lists': len(re.findall(r'^\d+\.', text, re.MULTILINE)),
            'bullet_points': len(re.findall(r'^[•\-\*]', text, re.MULTILINE)),
            'citations': len(re.findall(r'\[\d+\]|\(\d{4}\)', text)),
            'quotations': len(re.findall(r'"[^"]*"', text)),
            'parenthetical': len(re.findall(r'\([^)]*\)', text))
        }
        return features
    
    def _calculate_nonfiction_difficulty_score(self, common_ratio: float, academic_ratio: float,
                                             avg_sent_len: float, difficult_count: int, 
                                             unique_count: int, text_features: Dict) -> float:
        """计算非虚构文本难度评分"""
        # 基于研究文献的98%词汇覆盖率原则，但针对非虚构文本调整
        coverage_penalty = max(0, (0.95 - common_ratio) * 6)  # 非虚构文本允许更多专业词汇
        academic_bonus = min(academic_ratio * 2, 1.5)  # 适量学术词汇有助于理解
        sentence_penalty = max(0, (avg_sent_len - 18) * 0.12)  # 非虚构文本句子通常更长
        difficulty_penalty = min(difficult_count * 0.06, 2.0)
        
        # 文本特征调整 - 良好的结构化特征能降低难度
        structure_bonus = min(sum(text_features.values()) * 0.1, 1.0)
        
        base_score = 5
        total_adjustment = coverage_penalty + academic_bonus + sentence_penalty + difficulty_penalty - structure_bonus
        
        return min(10, max(1, base_score + total_adjustment))
    
    def _get_nonfiction_reading_level(self, score: float) -> str:
        """根据评分获取非虚构文本阅读水平"""
        if score <= 3:
            return "入门级 (适合非虚构文本初学者)"
        elif score <= 5:
            return "基础级 (适合有一定非虚构阅读经验者)"
        elif score <= 7:
            return "中级 (适合中等水平学术阅读者)"
        elif score <= 8.5:
            return "高级 (需要较强的学术阅读能力)"
        else:
            return "专业级 (需要专业领域知识背景)"
    
    def _estimate_nonfiction_reading_time(self, word_count: int) -> str:
        """估算非虚构文本阅读时间"""
        # 非虚构文本阅读速度通常比小说慢，需要更多思考时间
        minutes = word_count / 100  # 中国英语专业学生非虚构文本平均阅读速度
        
        if minutes < 1:
            return f"{int(minutes * 60)}秒"
        elif minutes < 60:
            return f"{int(minutes)}分钟"
        else:
            hours = int(minutes / 60)
            remaining_minutes = int(minutes % 60)
            return f"{hours}小时{remaining_minutes}分钟"

class EnhancedNonfictionReader:
    """增强版非虚构图书阅读助手"""
    
    def __init__(self, model_name: str = "huihui_ai/qwenlong-abliterated:latest"):
        self.model_name = model_name
        self.ollama_url = "http://localhost:11434/api/generate"
        self.processed_paragraphs = []
        self.difficulty_analyzer = TextDifficultyAnalyzer()
        self.vocab_db = VocabularyDatabase()
        # 可用模型列表
        self.available_models = [
            "huihui_ai/qwenlong-abliterated:latest",
            "gemma3:12b",
            "gemma3:27b", 
            "qwen3:32b",
            "qwen3:8b",
            "gemma3:4b",
            "phi4:latest"
        ]
    
    def set_model(self, model_name: str):
        """设置使用的模型"""
        if model_name in self.available_models:
            self.model_name = model_name
            logger.info(f"模型已切换为: {model_name}")
        else:
            logger.warning(f"模型 {model_name} 不在可用列表中")
    
    def create_enhanced_nonfiction_analysis_prompt(self, paragraph: str, difficulty_info: Dict) -> str:
        """创建增强的非虚构文本分析提示词（用于单段落详细分析）"""
        prompt = f"""
作为英语教学专家，请对以下英文非虚构文本段落进行深度分析，特别关注中国英语专业学生的学习需求：

【原文段落】
{paragraph}

【段落基本信息】
- 总词数：{difficulty_info['total_words']}
- 独特词汇：{difficulty_info['unique_words']}
- 词汇覆盖率：{difficulty_info['vocabulary_coverage']:.1f}%
- 学术词汇密度：{difficulty_info['academic_density']:.1f}%
- 难度等级：{difficulty_info['reading_level']}
- 预估阅读时间：{difficulty_info['estimated_reading_time']}
- 文本特征：{difficulty_info['text_features']}

请按照以下结构进行详细分析：

## 📊 非虚构文本难度评估
- 根据中国英语专业学生特点，评估此段落的学术阅读难度
- 分析文本结构特征对理解的影响（标题、列表、引用等）
- 指出可能造成理解障碍的语言特征

## 📚 核心词汇与术语深度解析
请选择5-8个关键词汇进行深度分析，重点关注：
- 学术词汇和专业术语的准确含义
- 词汇在特定学科语境中的用法
- 词族关系和词汇搭配
- 同义词、反义词和相关概念
- 在不同非虚构文本中的应用

## 🏗️ 论证结构与逻辑分析
- 识别文本的论证结构（因果、对比、分类等）
- 分析作者的论点、论据和论证方法
- 解释复杂句式结构和学术写作特征
- 识别信号词和连接词的逻辑作用

## 🎯 非虚构阅读策略指导
基于研究文献，提供具体的阅读策略：
- 预读策略：激活背景知识、预测内容
- 主动阅读技巧：标注、质疑、总结
- 文本特征利用：标题、副标题、视觉辅助
- 批判性思维：评估证据、识别偏见

## 🔍 深度内容分析（基于10大核心问题）
请结合以下关键分析维度深入探讨文本内容：

### 1️⃣ 核心问题识别
- 此段落试图解决或探讨的核心问题是什么？
- 作者在此段落中提出的主要论点或观点有哪些？

### 2️⃣ 证据与案例分析
- 作者提供了哪些重要证据、事实或案例来支持论点？
- 能否识别出关键的例证或数据？

### 3️⃣ 结构与逻辑顺序
- 此段落在整体论述中的位置和作用是什么？
- 段落内容如何围绕主题展开，呈现怎样的逻辑顺序？

### 4️⃣ 对立观点处理
- 作者是否在此段落中讨论或暗示相反的观点？
- 如何处理潜在的反对意见或争议？

### 5️⃣ 关键概念定义
- 段落中出现的关键概念或专业术语有哪些？
- 作者如何定义和解释这些概念？

### 6️⃣ 背景知识构建
- 作者提供了哪些背景知识或历史语境信息？
- 这些背景信息如何与主题相关联？

### 7️⃣ 实际应用价值
- 此段落提出了哪些实际建议、对策或结论？
- 这些观点在现实中有何意义或启示？

### 8️⃣ 独特见解识别
- 相较于该领域的其他观点，此段落有哪些独特之处？
- 作者的观点如何拓展读者对该领域的认识？

### 9️⃣ 写作风格分析
- 作者在此段落中的写作风格或论证方法有什么特点？
- 这种风格是否让内容更易理解或更具说服力？

### 🔟 核心启示提炼
- 此段落希望读者获得的最大收获或启示是什么？
- 对理解整本书的主题有何重要贡献？

## 🌍 背景知识与文化语境
- 提供必要的学科背景知识
- 解释文化、历史或社会语境
- 帮助理解作者的写作目的和受众
- 连接相关的概念框架

## 💡 批判性思考问题
基于10大核心问题框架，设计3-5个深层思考问题：
- 作者的论点是否有充分的证据支撑？存在哪些可能的反驳观点？
- 此段落的观点与该领域的其他理论或实践有何异同？
- 作者提供的背景信息是否足够帮助理解核心概念？
- 这些观点和建议在中国文化语境下是否同样适用？
- 阅读此段落后，你对该主题的理解发生了哪些变化？

## 🧠 理解检查与信息整合
- 主要论点和关键信息概括
- 论证逻辑和结构总结
- 理解程度自测问题
- 与其他相关知识的联系

## 📖 文本类型识别与特征分析
- 识别文本类型（学术文章、科普文章、传记等）
- 分析文本体裁特征和写作风格
- 说明该类型文本的阅读重点

## 🈶 精准中文翻译
提供两个版本的翻译：
1. 学术翻译版本（保持专业术语准确性）
2. 通俗理解版本（便于概念理解）

请确保分析深入、准确，特别关注中国英语专业学生在非虚构文本阅读中的具体需求和挑战。分析时要充分运用10大核心问题的分析框架，帮助学生建立系统性的非虚构文本理解能力。
"""
        return prompt
    
    def create_simplified_nonfiction_analysis_prompt(self, paragraph: str, difficulty_info: Dict) -> str:
        """创建简化的非虚构文本分析提示词（用于整本书处理）"""
        prompt = f"""
请对以下英文非虚构文本段落进行快速分析，为中国英语专业学生提供关键信息：

【原文段落】
{paragraph}

【段落信息】词数：{difficulty_info['total_words']}，学术密度：{difficulty_info['academic_density']:.1f}%，难度：{difficulty_info['reading_level']}

请提供简洁分析：

## 📚 关键术语（3-5个）
选择最重要的学术词汇或专业术语，简要说明含义和应用。

## 🏗️ 论证结构
简要说明文本的主要论点和论证逻辑。

## 🔍 核心内容要点（基于10大分析维度）
### 核心问题：此段落探讨的主要问题是什么？
### 关键证据：作者提供了哪些重要支撑材料？
### 逻辑结构：段落的组织逻辑和论述顺序如何？
### 概念定义：出现了哪些需要理解的关键概念？
### 实用价值：段落内容的现实意义和应用价值？

## 🎯 阅读要点
指出理解此段落的关键点和注意事项。

## 🈶 中文翻译
提供准确的学术翻译。

请保持简洁，重点突出核心学术内容和深度理解要素。
"""
        return prompt
    
    def call_ollama(self, prompt: str, is_simplified: bool = False) -> str:
        """调用ollama模型
        
        Args:
            prompt: 提示词
            is_simplified: 是否为简化分析（用于优化参数）
        """
        try:
            # 根据分析类型调整参数
            if is_simplified:
                # 简化分析：更低的温度，更少的tokens，更短的超时
                options = {
                    "temperature": 0.1,
                    "top_p": 0.8,
                    "max_tokens": 2000,
                    "repeat_penalty": 1.0,
                }
                timeout = 120  # 更短的超时时间
            else:
                # 详细分析：标准参数
                options = {
                    "temperature": 0.3,
                    "top_p": 0.9,
                    "max_tokens": 6000,
                    "repeat_penalty": 1.1,
                }
                timeout = 300
            
            payload = {
                "model": self.model_name,
                "prompt": prompt,
                "stream": False,
                "options": options
            }
            
            response = requests.post(self.ollama_url, json=payload, timeout=timeout)
            
            if response.status_code == 200:
                result = response.json()
                return result.get('response', '')
            else:
                logger.error(f"Ollama API error: {response.status_code}")
                return f"错误：API调用失败，状态码：{response.status_code}"
                
        except Exception as e:
            logger.error(f"Error calling Ollama: {str(e)}")
            return f"错误：{str(e)}"
    
    def analyze_paragraph(self, paragraph: str, index: int, use_detailed_analysis: bool = True) -> Dict[str, Any]:
        """分析段落
        
        Args:
            paragraph: 要分析的段落文本
            index: 段落索引
            use_detailed_analysis: 是否使用详细分析（True=详细分析，False=简化分析）
        """
        analysis_type = "详细" if use_detailed_analysis else "简化"
        logger.info(f"正在进行{analysis_type}分析第 {index + 1} 段落...")
        
        # 进行难度分析
        difficulty_info = self.difficulty_analyzer.analyze_text_difficulty(paragraph)
        
        # 根据分析类型选择提示词
        if use_detailed_analysis:
            prompt = self.create_enhanced_nonfiction_analysis_prompt(paragraph, difficulty_info)
        else:
            prompt = self.create_simplified_nonfiction_analysis_prompt(paragraph, difficulty_info)
        
        # 获取AI分析
        analysis = self.call_ollama(prompt, is_simplified=not use_detailed_analysis)
        
        # 提取并保存词汇
        self._extract_and_save_vocabulary(paragraph, analysis)
        
        result = {
            "index": index + 1,
            "original_text": paragraph,
            "difficulty_info": difficulty_info,
            "analysis": analysis,
            "analysis_type": analysis_type,
            "timestamp": time.strftime("%Y-%m-%d %H:%M:%S")
        }
        
        self.processed_paragraphs.append(result)
        return result
    
    def _extract_and_save_vocabulary(self, text: str, analysis: str):
        """从文本和分析中提取词汇并保存到数据库"""
        # 简单的词汇提取（可以后续改进）
        words = word_tokenize(text.lower())
        words = [word for word in words if word.isalpha() and len(word) > 3]
        
        for word in set(words):
            if word not in self.difficulty_analyzer.common_words:
                # 这里可以添加更复杂的词汇定义提取逻辑
                self.vocab_db.add_word(word, "", "", 5)
    
    def get_nonfiction_reading_recommendations(self, difficulty_score: float) -> List[str]:
        """根据难度评分提供非虚构文本阅读建议"""
        recommendations = []
        
        if difficulty_score > 8:
            recommendations.extend([
                "🚨 此非虚构文本难度较高，建议：",
                "• 先预习相关学科背景知识和专业术语",
                "• 采用SQ3R阅读法：浏览、质疑、阅读、复述、复习",
                "• 重点关注文本结构和论证逻辑",
                "• 使用学术词典和专业资源辅助理解",
                "• 做好详细笔记和概念图"
            ])
        elif difficulty_score > 6:
            recommendations.extend([
                "⚠️ 此非虚构文本具有一定学术挑战性，建议：",
                "• 预读时重点关注标题、副标题和文本特征",
                "• 识别主要论点和支撑证据",
                "• 积极运用批判性思维评估信息",
                "• 联系已有知识构建理解框架",
                "• 适当查阅背景资料"
            ])
        else:
            recommendations.extend([
                "✅ 此非虚构文本难度适中，建议：",
                "• 保持主动阅读，边读边思考",
                "• 注意文本的组织结构和逻辑关系",
                "• 练习总结和概括关键信息",
                "• 思考作者观点与你的观点差异",
                "• 享受学习新知识的过程"
            ])
        
        return recommendations
    
    def split_text_into_sections(self, text: str) -> List[str]:
        """智能分割非虚构文本为段落或章节"""
        # 先尝试按章节分割
        sections = re.split(r'\n\s*(?:Chapter|Section|Part|\d+\.)\s+[A-Z].*?\n', text.strip())
        
        if len(sections) <= 1:
            # 如果没有明显章节，按段落分割
            sections = re.split(r'\n\s*\n', text.strip())
        
        cleaned_sections = []
        for section in sections:
            section = section.strip()
            if section and len(section.split()) >= 30:  # 非虚构文本段落通常更长
                cleaned_sections.append(section)
        
        return cleaned_sections
    
    def create_enhanced_nonfiction_docx(self, book_title: str = "英文非虚构图书阅读分析报告") -> str:
        """创建增强版非虚构文本DOCX文档"""
        doc = Document()
        
        title = doc.add_heading('📖 ' + book_title, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph(f"生成时间：{time.strftime('%Y年%m月%d日 %H:%M:%S')}")
        doc.add_paragraph(f"处理段落数：{len(self.processed_paragraphs)}")
        doc.add_paragraph("基于《How to Read Non-Fiction English Books for Chinese English Majors》研究文献")
        doc.add_paragraph("采用10大核心问题深度分析框架：核心问题识别、证据案例分析、结构逻辑梳理、对立观点处理、关键概念定义、背景知识构建、实际应用价值、独特见解识别、写作风格分析、核心启示提炼")
        
        # 添加分析模式信息
        if self.processed_paragraphs:
            analysis_mode = self.processed_paragraphs[0].get('analysis_type', '详细')
            doc.add_paragraph(f"分析模式：{analysis_mode}分析")
            if analysis_mode == "简化":
                doc.add_paragraph("⚡ 采用快速分析模式，专注核心学术内容")
        
        doc.add_paragraph("=" * 60)
        
        if self.processed_paragraphs:
            # 添加总体统计
            doc.add_heading('📊 非虚构文本阅读统计概览', level=1)
            
            total_words = sum(p['difficulty_info']['total_words'] for p in self.processed_paragraphs)
            avg_difficulty = sum(p['difficulty_info']['difficulty_score'] for p in self.processed_paragraphs) / len(self.processed_paragraphs)
            avg_academic_density = sum(p['difficulty_info']['academic_density'] for p in self.processed_paragraphs) / len(self.processed_paragraphs)
            
            doc.add_paragraph(f"• 总词数：{total_words}")
            doc.add_paragraph(f"• 平均难度评分：{avg_difficulty:.1f}/10")
            doc.add_paragraph(f"• 平均学术词汇密度：{avg_academic_density:.1f}%")
            doc.add_paragraph(f"• 预估总阅读时间：{self.difficulty_analyzer._estimate_nonfiction_reading_time(total_words)}")
            
            # 添加每个段落的分析
            for paragraph_data in self.processed_paragraphs:
                doc.add_page_break()
                
                doc.add_heading(f"第 {paragraph_data['index']} 段", level=1)
                
                difficulty_info = paragraph_data['difficulty_info']
                doc.add_heading('📊 难度评估', level=2)
                doc.add_paragraph(f"难度评分：{difficulty_info['difficulty_score']:.1f}/10")
                doc.add_paragraph(f"阅读等级：{difficulty_info['reading_level']}")
                doc.add_paragraph(f"词汇覆盖率：{difficulty_info['vocabulary_coverage']:.1f}%")
                doc.add_paragraph(f"学术词汇密度：{difficulty_info['academic_density']:.1f}%")
                
                doc.add_heading('📖 原文', level=2)
                p = doc.add_paragraph(paragraph_data['original_text'])
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                
                doc.add_heading('🔍 详细分析', level=2)
                analysis_p = doc.add_paragraph(paragraph_data['analysis'])
                analysis_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                
                recommendations = self.get_nonfiction_reading_recommendations(difficulty_info['difficulty_score'])
                if recommendations:
                    doc.add_heading('💡 阅读建议', level=2)
                    for rec in recommendations:
                        doc.add_paragraph(rec)
        
        filename = f"nonfiction_analysis_{time.strftime('%Y%m%d_%H%M%S')}.docx"
        doc.save(filename)
        return filename

class EnhancedGradioInterface:
    """增强版非虚构图书阅读界面"""
    
    def __init__(self):
        self.reader = EnhancedNonfictionReader()
        self.current_paragraphs = []
        self.current_index = 0
        self.current_book_title = "未命名非虚构图书"
        self.current_model = self.reader.model_name
    
    def change_model(self, model_name: str) -> str:
        """切换模型"""
        self.reader.set_model(model_name)
        self.current_model = model_name
        return f"✅ 已切换到模型：{model_name}"
        
    def handle_file_upload(self, uploaded_file_path) -> Tuple[str, str]:
        """处理上传的文件"""
        try:
            if uploaded_file_path is None or uploaded_file_path == "":
                return "❌ 请选择要上传的文件", ""
            
            # 检查文件类型
            if not uploaded_file_path.lower().endswith(('.txt', '.md')):
                return "❌ 只支持 .txt 和 .md 格式的文件", ""
            
            # 获取文件名
            self.current_book_title = os.path.splitext(os.path.basename(uploaded_file_path))[0]
            
            # 读取文件内容
            with open(uploaded_file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            return self._load_content(content)
            
        except Exception as e:
            return f"❌ 上传文件时出错：{str(e)}", ""
    
    def load_and_analyze_book(self, file_path: str) -> Tuple[str, str]:
        """加载并分析非虚构图书"""
        try:
            if file_path and os.path.exists(file_path):
                self.current_book_title = os.path.splitext(os.path.basename(file_path))[0]
                
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
                
                return self._load_content(content)
            else:
                return "❌ 文件路径无效或文件不存在", ""
        except Exception as e:
            return f"❌ 加载文件时出错：{str(e)}", ""
    
    def handle_text_input(self, text_input: str) -> Tuple[str, str]:
        """处理用户直接输入的文本"""
        try:
            if not text_input or not text_input.strip():
                return "❌ 请输入要分析的文本内容", ""
            
            # 设置标题
            self.current_book_title = "用户输入文本"
            
            return self._load_content(text_input.strip())
            
        except Exception as e:
            return f"❌ 处理文本输入时出错：{str(e)}", ""
    
    def analyze_single_text(self, text_input: str) -> Tuple[str, str, str]:
        """分析用户输入的单段文本"""
        try:
            if not text_input or not text_input.strip():
                return "❌ 请输入要分析的文本内容", "", ""
            
            text = text_input.strip()
            
            # 重置状态
            self.reader.processed_paragraphs = []
            
            # 分析文本
            result = self.reader.analyze_paragraph(text, 0, use_detailed_analysis=True)
            
            difficulty_info = result['difficulty_info']
            difficulty_display = f"""📊 文本难度分析：
• 文本类型：非虚构文本段落
• 难度评分：{difficulty_info['difficulty_score']:.1f}/10
• 阅读等级：{difficulty_info['reading_level']}
• 词汇覆盖率：{difficulty_info['vocabulary_coverage']:.1f}%
• 学术词汇密度：{difficulty_info['academic_density']:.1f}%
• 预估阅读时间：{difficulty_info['estimated_reading_time']}
• 总词数：{difficulty_info['total_words']}，独特词汇：{difficulty_info['unique_words']}
• 专业术语数量：{len(difficulty_info['technical_terms'])}
• 文本特征：{difficulty_info['text_features']}"""
            
            progress_info = "✅ 单段文本分析完成"
            
            return progress_info, difficulty_display, result['analysis']
            
        except Exception as e:
            error_message = f"❌ 分析文本时出现错误：{str(e)}"
            logger.error(error_message)
            return error_message, "", ""
    
    def _load_content(self, content: str) -> Tuple[str, str]:
        """加载内容的共同逻辑"""
        self.current_paragraphs = self.reader.split_text_into_sections(content)
        self.current_index = 0
        self.reader.processed_paragraphs = []
        
        overall_difficulty = self.reader.difficulty_analyzer.analyze_text_difficulty(content)
        
        status_message = f"✅ 成功加载非虚构图书《{self.current_book_title}》，共 {len(self.current_paragraphs)} 段落"
        
        difficulty_summary = f"""📊 整体难度分析：
• 总词数：{overall_difficulty['total_words']:,}
• 独特词汇：{overall_difficulty['unique_words']:,}
• 词汇覆盖率：{overall_difficulty['vocabulary_coverage']:.1f}%
• 学术词汇密度：{overall_difficulty['academic_density']:.1f}%
• 难度评分：{overall_difficulty['difficulty_score']:.1f}/10
• 阅读等级：{overall_difficulty['reading_level']}
• 预估阅读时间：{overall_difficulty['estimated_reading_time']}
• 专业术语数量：{len(overall_difficulty.get('technical_terms', []))}
• 文本特征统计：{overall_difficulty.get('text_features', {})}

💡 非虚构阅读建议：
{chr(10).join(self.reader.get_nonfiction_reading_recommendations(overall_difficulty['difficulty_score']))}"""
        
        return status_message, difficulty_summary
    
    def process_next_paragraph(self) -> Tuple[str, str, str, str]:
        """处理下一个段落"""
        if not self.current_paragraphs:
            return "❌ 请先加载非虚构图书文件", "", "", ""
        
        if self.current_index >= len(self.current_paragraphs):
            return "✅ 所有段落已处理完成", "", "", ""
        
        current_paragraph = self.current_paragraphs[self.current_index]
        result = self.reader.analyze_paragraph(current_paragraph, self.current_index)
        self.current_index += 1
        
        progress_info = f"已处理 {self.current_index}/{len(self.current_paragraphs)} 段落"
        
        difficulty_info = result['difficulty_info']
        difficulty_display = f"""📊 当前段落难度：
• 难度评分：{difficulty_info['difficulty_score']:.1f}/10
• 阅读等级：{difficulty_info['reading_level']}
• 词汇覆盖率：{difficulty_info['vocabulary_coverage']:.1f}%
• 预估阅读时间：{difficulty_info['estimated_reading_time']}
• 总词数：{difficulty_info['total_words']}，独特词汇：{difficulty_info['unique_words']}"""
        
        return progress_info, difficulty_display, result['original_text'], result['analysis']
    
    def process_entire_book(self) -> str:
        """处理整本非虚构图书（使用简化分析模式，专注核心学术内容）"""
        if not self.current_paragraphs:
            return "❌ 请先加载非虚构图书文件"
        
        try:
            total_paragraphs = len(self.current_paragraphs)
            logger.info(f"开始处理整本非虚构图书《{self.current_book_title}》，共 {total_paragraphs} 段落")
            logger.info("📈 使用简化分析模式，专注核心学术内容")
            
            # 重置处理状态
            self.reader.processed_paragraphs = []
            
            # 处理所有段落 - 使用简化分析模式
            for i, paragraph in enumerate(self.current_paragraphs):
                logger.info(f"正在快速处理第 {i+1}/{total_paragraphs} 段落")
                # 使用 use_detailed_analysis=False 来使用简化模式
                result = self.reader.analyze_paragraph(paragraph, i, use_detailed_analysis=False)
            
            # 保存完整分析
            filename = self.reader.create_enhanced_nonfiction_docx(self.current_book_title)
            
            final_message = f"""✅ 整本非虚构图书快速处理完成！

📖 非虚构图书名称：《{self.current_book_title}》
📊 处理统计：共处理 {total_paragraphs} 个段落
📄 分析报告：已保存为 {filename}
⚡ 分析模式：简化模式（专注核心学术内容）

🎯 分析包含：
• 每段落的关键术语和学术词汇分析
• 重要论证结构和逻辑分析
• 学术阅读策略指导
• 精准的学术翻译

📚 您可以打开 {filename} 查看完整的分析报告！

💡 提示：如需详细分析，请使用"处理下一段"功能逐段分析。"""
            
            logger.info(f"整本非虚构图书快速处理完成，报告保存为：{filename}")
            return final_message
            
        except Exception as e:
            error_message = f"❌ 处理过程中出现错误：{str(e)}"
            logger.error(error_message)
            return error_message
    
    def save_enhanced_analysis(self) -> str:
        """保存增强分析结果"""
        if not self.reader.processed_paragraphs:
            return "❌ 没有已处理的段落可以保存"
        
        filename = self.reader.create_enhanced_nonfiction_docx(self.current_book_title)
        return f"✅ 增强分析报告已保存为 {filename}，共包含 {len(self.reader.processed_paragraphs)} 个段落的详细分析"

def create_enhanced_interface():
    """创建增强版Gradio界面"""
    interface = EnhancedGradioInterface()
    
    with gr.Blocks(title="英文非虚构图书阅读辅助软件", theme=gr.themes.Soft()) as demo:
        gr.Markdown("# 📖 英文非虚构图书阅读辅助软件")
        gr.Markdown("**专为中国英语专业学生设计的非虚构文本阅读助手**")
        gr.Markdown("Designed by Toby")
        
        gr.Markdown("""
        ## 🔍 核心分析框架：10大深度问题
        
        **本软件基于非虚构文本阅读的10大核心分析维度：**
        
        1. **核心问题识别** - 探讨文本要解决的核心问题和主要论点
        2. **证据案例分析** - 识别关键证据、事实和案例研究  
        3. **结构逻辑梳理** - 分析文本组织结构和逻辑顺序
        4. **对立观点处理** - 识别和分析反对观点及其反驳
        5. **关键概念定义** - 解析专业术语和核心概念
        6. **背景知识构建** - 提供历史语境和背景信息
        7. **实际应用价值** - 探讨现实意义和实用建议
        8. **独特见解识别** - 发现创新观点和独特贡献
        9. **写作风格分析** - 分析论证方法和表达特色
        10. **核心启示提炼** - 总结主要收获和深层启示
        
        通过这一系统性分析框架，帮助学生建立深度的非虚构文本理解能力。
        """)
        
        with gr.Row():
            with gr.Column(scale=1):
                gr.Markdown("## 🤖 模型设置")
                model_dropdown = gr.Dropdown(
                    choices=interface.reader.available_models,
                    value=interface.current_model,
                    label="选择AI模型",
                    info="不同模型有不同的特点，可根据需要切换"
                )
                model_status = gr.Textbox(label="模型状态", interactive=False)
                
                gr.Markdown("## 📁 文本来源")
                
                # 添加多种文本输入方式
                with gr.Tabs():
                    with gr.TabItem("✍️ 直接输入文本"):
                        text_input = gr.Textbox(
                            label="输入要分析的非虚构文本",
                            placeholder="在此粘贴或输入英文非虚构文本内容...\n\n支持学术文章、科普文章、传记、历史文本等各类非虚构文本",
                            lines=8,
                            max_lines=15
                        )
                        with gr.Row():
                            analyze_text_btn = gr.Button("🔍 分析此文本", variant="primary")
                            load_text_btn = gr.Button("📖 加载为图书", variant="secondary")
                    
                    with gr.TabItem("📁 上传文件"):
                        file_upload = gr.File(
                            label="选择非虚构图书文件 (.txt 或 .md)",
                            file_types=['.txt', '.md'],
                            type="filepath"
                        )
                        upload_btn = gr.Button("📤 上传并分析", variant="primary")
                    
                    with gr.TabItem("📂 本地文件"):
                        file_input = gr.Textbox(
                            label="非虚构图书文件路径",
                            placeholder="输入非虚构图书文件的完整路径...",
                            value=""
                        )
                        load_btn = gr.Button("📊 加载并分析", variant="primary")
                
                load_status = gr.Textbox(label="加载状态", interactive=False)
                
                difficulty_analysis = gr.Textbox(
                    label="📊 难度分析报告", 
                    lines=12, 
                    interactive=False,
                    placeholder="文本难度分析将在这里显示..."
                )
                
                gr.Markdown("## 🎯 处理选项")
                gr.Markdown("**📝 处理模式说明：**")
                gr.Markdown("• **详细分析模式**：深度分析，包含完整的学术指导，适合学习研究")
                gr.Markdown("• **快速处理模式**：高效分析整本图书，快速生成报告")
                
                progress_info = gr.Textbox(label="处理进度", interactive=False)
                
                with gr.Row():
                    next_btn = gr.Button("➡️ 处理下一段（详细模式）", variant="secondary")
                    process_all_btn = gr.Button("🚀 处理整本图书（快速模式）", variant="primary")
                
                with gr.Row():
                    save_btn = gr.Button("💾 保存当前分析", variant="secondary")
        
        with gr.Row():
            with gr.Column(scale=2):
                current_difficulty = gr.Textbox(
                    label="📊 当前段落难度信息",
                    lines=8,
                    interactive=False,
                    placeholder="段落难度信息将在这里显示..."
                )
                
                gr.Markdown("## 📖 英文原文")
                original_text = gr.Textbox(
                    label="原文内容",
                    lines=10,
                    interactive=False,
                    placeholder="英文原文将在这里显示..."
                )
                
                gr.Markdown("## 🔍 深度分析结果")
                analysis_result = gr.Textbox(
                    label="专业分析",
                    lines=25,
                    interactive=False,
                    placeholder="详细的非虚构文本分析结果将在这里显示..."
                )
        
        # 事件绑定
        model_dropdown.change(
            fn=interface.change_model,
            inputs=[model_dropdown],
            outputs=[model_status]
        )
        
        # 直接文本分析
        analyze_text_btn.click(
            fn=interface.analyze_single_text,
            inputs=[text_input],
            outputs=[progress_info, current_difficulty, analysis_result]
        )
        
        # 加载文本为图书
        load_text_btn.click(
            fn=interface.handle_text_input,
            inputs=[text_input],
            outputs=[load_status, difficulty_analysis]
        )
        
        upload_btn.click(
            fn=interface.handle_file_upload,
            inputs=[file_upload],
            outputs=[load_status, difficulty_analysis]
        )
        
        load_btn.click(
            fn=interface.load_and_analyze_book,
            inputs=[file_input],
            outputs=[load_status, difficulty_analysis]
        )
        
        next_btn.click(
            fn=interface.process_next_paragraph,
            inputs=[],
            outputs=[progress_info, current_difficulty, original_text, analysis_result]
        )
        
        process_all_btn.click(
            fn=interface.process_entire_book,
            inputs=[],
            outputs=[progress_info]
        )
        
        save_btn.click(
            fn=interface.save_enhanced_analysis,
            inputs=[],
            outputs=[progress_info]
        )
    
    return demo

if __name__ == "__main__":
    demo = create_enhanced_interface()
    demo.launch(
        server_name="127.0.0.1",
        server_port=7865,
        share=False,
        show_error=True
    ) 